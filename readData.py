from docx.api import Document

class Docs:
    
    def __init__(self):

        self.document = Document("cse_courses.docx")


# Extract text from tables

    def process_table_2(self, courses):
        
        document = self.document
        
        
        for row in document.tables[2].rows[2:]:
                
                course_code = row.cells[1].text.strip()
                course_name = row.cells[2].text.strip()
                
                course_code = course_code.split()

                

                course_code = "".join(course_code)
                
                            
                courses.append((course_code, course_name))
                
        return courses
    
    def process_table_1(self, courses):
        
        document = self.document
        
        table_1 = document.tables[1]
        
        for row in table_1.rows[2:10]:
            
            course_code = row.cells[1].text.strip()
            course_name = row.cells[2].text.strip()
            
            course_code = course_code.split()

            

            course_code = "".join(course_code)
            
                        
            courses.append((course_code, course_name))
            
            
        for row in table_1.rows[-4:]:
            
            course_code = row.cells[1].text.strip()
            course_name = row.cells[2].text.strip()
            
            course_code = course_code.split()

            

            course_code = "".join(course_code)
            
                        
            courses.append((course_code, course_name))
                
                
        return courses
            

    def get_courses(self):
        courses = []
        courses = self.process_table_1(courses)
        courses = self.process_table_2(courses)
        
        return courses
        
